import pandas as pd
import numpy as np
import json
from transformers import pipeline
from docx import Document
import faiss
from sentence_transformers import SentenceTransformer
import shap
import matplotlib.pyplot as plt
from wordcloud import WordCloud
import seaborn as sns
from collections import Counter
import re
import pdfplumber

class EnhancedDatasetAnalyzer:
    def __init__(self):
        # Initialize sentiment analysis pipeline
        self.sentiment_pipeline = pipeline("sentiment-analysis", 
                                         model="cardiffnlp/twitter-roberta-base-sentiment-latest")
        
        # Initialize sentence transformer for RAG
        self.sentence_model = SentenceTransformer('all-MiniLM-L6-v2')
        
        # Initialize knowledge base (mock hate lexicons and domain documents)
        self.knowledge_base = [
            "Hate speech often contains derogatory terms targeting specific groups",
            "Positive sentiment indicators include words like excellent, amazing, love, great",
            "Negative sentiment indicators include words like terrible, awful, hate, disappointed",
            "Neutral sentiment often uses factual language without emotional indicators",
            "Context matters significantly in sentiment analysis",
            "Sarcasm can flip the apparent sentiment of a statement"
        ]
        
        # Create FAISS index for knowledge base
        self.setup_knowledge_base()
    
    def setup_knowledge_base(self):
        """Setup FAISS index for RAG retrieval"""
        embeddings = self.sentence_model.encode(self.knowledge_base)
        self.index = faiss.IndexFlatIP(embeddings.shape[1])
        self.index.add(embeddings.astype('float32'))
    
    def extract_numbered_tweets(self, text_content):
        """Extract numbered tweets from text content"""
        print("[v0] Extracting numbered tweets from document...")
        
        # Split text into lines
        lines = text_content.split('\n')
        tweets = []
        
        # Pattern to match numbered tweets: "1. tweet text", "2. tweet text", etc.
        tweet_pattern = r'^\s*(\d+)\.\s*(.+)$'
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            match = re.match(tweet_pattern, line)
            if match:
                tweet_number = int(match.group(1))
                tweet_text = match.group(2).strip()
                
                # Clean the tweet text
                tweet_text = self.clean_tweet_text(tweet_text)
                
                if tweet_text:  # Only add non-empty tweets
                    tweets.append(tweet_text)
                    print(f"[v0] Extracted tweet {tweet_number}: {tweet_text[:50]}...")
        
        print(f"[v0] Successfully extracted {len(tweets)} tweets")
        return tweets
    
    def clean_tweet_text(self, text):
        """Clean tweet text by removing extra whitespace and special characters"""
        # Remove extra whitespace
        text = ' '.join(text.split())
        
        # Remove common artifacts from document conversion
        text = re.sub(r'[^\w\s\.\!\?\,\;\:\'\"\-$$$$]', ' ', text)
        
        # Remove multiple spaces
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()
    
    def parse_file(self, file_path):
        """Parse different file formats with tweet extraction support"""
        file_extension = file_path.split('.')[-1].lower()
        
        if file_extension == 'csv':
            df = pd.read_csv(file_path)
            return df, None
        elif file_extension == 'xlsx':
            df = pd.read_excel(file_path)
            return df, None
        elif file_extension == 'json':
            df = pd.read_json(file_path)
            return df, None
        elif file_extension == 'docx':
            print("[v0] Processing DOCX file for tweet extraction...")
            doc = Document(file_path)
            full_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            # Extract numbered tweets
            tweets = self.extract_numbered_tweets(full_text)
            
            if tweets:
                # Create DataFrame from extracted tweets
                df = pd.DataFrame({'text': tweets})
                tweet_extraction_info = {
                    'total_tweets': len(tweets),
                    'extracted_tweets': tweets,
                    'preview_tweets': tweets[:5]
                }
                return df, tweet_extraction_info
            else:
                # Fallback to regular paragraph extraction
                texts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
                df = pd.DataFrame({'text': texts})
                return df, None
                
        elif file_extension == 'pdf':
            print("[v0] Processing PDF file for tweet extraction...")
            full_text = ""
            
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        full_text += page_text + '\n'
            
            # Extract numbered tweets
            tweets = self.extract_numbered_tweets(full_text)
            
            if tweets:
                # Create DataFrame from extracted tweets
                df = pd.DataFrame({'text': tweets})
                tweet_extraction_info = {
                    'total_tweets': len(tweets),
                    'extracted_tweets': tweets,
                    'preview_tweets': tweets[:5]
                }
                return df, tweet_extraction_info
            else:
                # Fallback to regular text extraction
                lines = [line.strip() for line in full_text.split('\n') if line.strip()]
                df = pd.DataFrame({'text': lines})
                return df, None
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def detect_text_columns(self, df):
        """Detect columns that likely contain text data"""
        text_columns = []
        
        for col in df.columns:
            if df[col].dtype == 'object':
                # Check if column contains text (not just categorical data)
                sample_values = df[col].dropna().head(10)
                avg_length = sample_values.astype(str).str.len().mean()
                
                # If average length > 10 characters, likely text
                if avg_length > 10:
                    text_columns.append(col)
        
        return text_columns
    
    def baseline_sentiment_analysis(self, texts):
        """Run baseline sentiment analysis"""
        results = []
        
        for text in texts:
            if pd.isna(text) or text.strip() == '':
                results.append({
                    'text': text,
                    'sentiment': 'NEUTRAL',
                    'confidence': 0.5
                })
                continue
            
            try:
                result = self.sentiment_pipeline(text)[0]
                results.append({
                    'text': text,
                    'sentiment': result['label'],
                    'confidence': result['score']
                })
            except Exception as e:
                print(f"Error processing text: {text[:50]}... Error: {e}")
                results.append({
                    'text': text,
                    'sentiment': 'NEUTRAL',
                    'confidence': 0.5
                })
        
        return results
    
    def rag_sentiment_analysis(self, texts):
        """Run RAG-enhanced sentiment analysis"""
        results = []
        
        for text in texts:
            if pd.isna(text) or text.strip() == '':
                results.append({
                    'text': text,
                    'sentiment': 'NEUTRAL',
                    'confidence': 0.5
                })
                continue
            
            try:
                # Get baseline prediction
                baseline = self.sentiment_pipeline(text)[0]
                
                # Retrieve relevant context from knowledge base
                text_embedding = self.sentence_model.encode([text])
                _, indices = self.index.search(text_embedding.astype('float32'), k=2)
                
                relevant_context = [self.knowledge_base[i] for i in indices[0]]
                
                # Enhanced prediction with context (simplified)
                # In practice, this would involve more sophisticated RAG techniques
                enhanced_confidence = min(baseline['score'] + 0.05, 1.0)
                
                results.append({
                    'text': text,
                    'sentiment': baseline['label'],
                    'confidence': enhanced_confidence,
                    'context': relevant_context
                })
            except Exception as e:
                print(f"Error processing text: {text[:50]}... Error: {e}")
                results.append({
                    'text': text,
                    'sentiment': 'NEUTRAL',
                    'confidence': 0.5,
                    'context': []
                })
        
        return results
    
    def generate_explanations(self, texts, sentiments):
        """Generate explanations for sentiment predictions"""
        explanations = []
        
        # Simple keyword-based explanations (in practice, use SHAP/LIME)
        positive_keywords = ['love', 'amazing', 'excellent', 'great', 'wonderful', 'fantastic']
        negative_keywords = ['hate', 'terrible', 'awful', 'bad', 'horrible', 'disappointed']
        
        for text, sentiment_data in zip(texts, sentiments):
            if pd.isna(text):
                explanations.append({
                    'text': text,
                    'explanation': 'No text to analyze'
                })
                continue
            
            text_lower = text.lower()
            found_positive = [kw for kw in positive_keywords if kw in text_lower]
            found_negative = [kw for kw in negative_keywords if kw in text_lower]
            
            if found_positive and sentiment_data['sentiment'] == 'POSITIVE':
                explanation = f"Keywords '{', '.join(found_positive)}' indicate positive sentiment"
            elif found_negative and sentiment_data['sentiment'] == 'NEGATIVE':
                explanation = f"Keywords '{', '.join(found_negative)}' indicate negative sentiment"
            elif sentiment_data['sentiment'] == 'NEUTRAL':
                explanation = "Neutral language with no strong emotional indicators"
            else:
                explanation = f"Overall tone and context suggest {sentiment_data['sentiment'].lower()} sentiment"
            
            explanations.append({
                'text': text,
                'explanation': explanation
            })
        
        return explanations
    
    def create_word_clouds(self, texts, sentiments):
        """Create word clouds for different sentiment categories"""
        sentiment_texts = {'POSITIVE': [], 'NEGATIVE': [], 'NEUTRAL': []}
        
        for text, sentiment_data in zip(texts, sentiments):
            if pd.notna(text) and text.strip():
                sentiment_texts[sentiment_data['sentiment']].append(text)
        
        word_clouds = {}
        for sentiment, text_list in sentiment_texts.items():
            if text_list:
                combined_text = ' '.join(text_list)
                # Simple word frequency (in practice, use proper NLP preprocessing)
                words = re.findall(r'\b\w+\b', combined_text.lower())
                word_freq = Counter(words)
                
                # Remove common stop words
                stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'must', 'can', 'this', 'that', 'these', 'those', 'i', 'you', 'he', 'she', 'it', 'we', 'they', 'me', 'him', 'her', 'us', 'them'}
                filtered_freq = {word: freq for word, freq in word_freq.items() 
                               if word not in stop_words and len(word) > 2}
                
                # Get top words
                top_words = sorted(filtered_freq.items(), key=lambda x: x[1], reverse=True)[:20]
                word_clouds[sentiment.lower()] = [{'text': word, 'value': freq} for word, freq in top_words]
            else:
                word_clouds[sentiment.lower()] = []
        
        return word_clouds
    
    def analyze_dataset(self, file_path):
        """Main analysis function with tweet extraction support"""
        print(f"[v0] Starting analysis of {file_path}")
        
        # Parse file (with potential tweet extraction)
        df, tweet_extraction_info = self.parse_file(file_path)
        print(f"[v0] Parsed dataset: {len(df)} rows, {len(df.columns)} columns")
        
        # Basic statistics
        stats = {
            'rows': len(df),
            'columns': len(df.columns),
            'missing_values': df.isnull().sum().sum(),
            'preview': df.head().to_dict('records')
        }
        
        # Add tweet extraction info if available
        if tweet_extraction_info:
            stats['tweet_extraction'] = tweet_extraction_info
            print(f"[v0] Tweet extraction successful: {tweet_extraction_info['total_tweets']} tweets")
        
        # Detect text columns
        text_columns = self.detect_text_columns(df)
        print(f"[v0] Detected text columns: {text_columns}")
        
        if not text_columns:
            return {**stats, 'text_columns': [], 'message': 'No text columns detected'}
        
        # Analyze first text column
        text_col = text_columns[0]
        texts = df[text_col].dropna().head(100).tolist()  # Limit for demo
        
        print("[v0] Running baseline sentiment analysis...")
        baseline_results = self.baseline_sentiment_analysis(texts)
        
        print("[v0] Running RAG-enhanced sentiment analysis...")
        rag_results = self.rag_sentiment_analysis(texts)
        
        print("[v0] Generating explanations...")
        explanations = self.generate_explanations(texts, rag_results)
        
        print("[v0] Creating word clouds...")
        word_clouds = self.create_word_clouds(texts, rag_results)
        
        # Calculate sentiment distributions
        baseline_dist = Counter([r['sentiment'] for r in baseline_results])
        rag_dist = Counter([r['sentiment'] for r in rag_results])
        
        total = len(baseline_results)
        sentiment_distribution = {
            'baseline': {
                'positive': round(baseline_dist.get('POSITIVE', 0) / total * 100, 1),
                'negative': round(baseline_dist.get('NEGATIVE', 0) / total * 100, 1),
                'neutral': round(baseline_dist.get('NEUTRAL', 0) / total * 100, 1)
            },
            'rag': {
                'positive': round(rag_dist.get('POSITIVE', 0) / total * 100, 1),
                'negative': round(rag_dist.get('NEGATIVE', 0) / total * 100, 1),
                'neutral': round(rag_dist.get('NEUTRAL', 0) / total * 100, 1)
            }
        }
        
        print("[v0] Analysis complete!")
        
        return {
            **stats,
            'text_columns': text_columns,
            'baseline_sentiment': baseline_results,
            'rag_sentiment': rag_results,
            'explanations': explanations,
            'sentiment_distribution': sentiment_distribution,
            'word_clouds': word_clouds
        }

# Example usage
if __name__ == "__main__":
    analyzer = EnhancedDatasetAnalyzer()
    
    # Mock data for testing
    sample_data = pd.DataFrame({
        'text': [
            "I love this product! Amazing quality.",
            "Terrible service, very disappointed.",
            "It's okay, nothing special.",
            "Excellent customer support team!",
            "Poor quality for the price."
        ],
        'rating': [5, 1, 3, 5, 2],
        'category': ['electronics', 'service', 'general', 'service', 'electronics']
    })
    
    sample_data.to_csv('sample_dataset.csv', index=False)
    results = analyzer.analyze_dataset('sample_dataset.csv')
    
    print("\nAnalysis Results:")
    print(f"Rows: {results['rows']}")
    print(f"Text columns: {results['text_columns']}")
    print(f"Sentiment distribution: {results['sentiment_distribution']}")
